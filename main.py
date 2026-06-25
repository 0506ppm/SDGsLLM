# ✅ main.py（TWCC）- 支援 URL 下載 + 優化文字解析
from fastapi import FastAPI, File, UploadFile
from pydantic import BaseModel
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, AutoModel
from peft import PeftModel, PeftConfig
from pyngrok import ngrok, conf
import torch
import faiss
import json
import os
import uuid
import shutil
import numpy as np
from fastapi.middleware.cors import CORSMiddleware
import fitz  # PyMuPDF
import docx
from tika import parser
import tika
import requests
from dotenv import load_dotenv
import re  # 新增：用於文字清理
from groq import Groq  # 新增：Groq API
import cohere

load_dotenv()

# 初始化 Tika
tika.initVM()

local_api = os.getenv("LOCAL_API")
groq_api_key = os.getenv("GROQ_API_KEY")
cohere_api_key = os.getenv("COHERE_API_KEY")

# 初始化 Groq 客戶端
groq_client = Groq(api_key=groq_api_key)

# 初始化 Cohere 客戶端
if cohere_api_key:
    try:
        cohere_client = cohere.Client(cohere_api_key)
        print("✅ Cohere Rerank 已啟用")
    except Exception as e:
        cohere_client = None
        print(f"⚠️ Cohere 初始化失敗: {e}")
else:
    cohere_client = None
    print("⚠️ COHERE_API_KEY 未設定，Rerank 功能將被停用")

# === Ngrok 設定 ===
conf.get_default().auth_token = os.getenv("NGROK_TOKEN")
app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# === 載入 RAG 模型 ===
print("\U0001F680 載入 LLM 模型中...")
lora_model_path = "./output/result_mistral7b-lora"
base_model_path = PeftConfig.from_pretrained(lora_model_path).base_model_name_or_path

tokenizer = AutoTokenizer.from_pretrained(base_model_path)
base_model = AutoModelForCausalLM.from_pretrained(
    base_model_path,
    device_map="auto",
    offload_folder="./offload",
    torch_dtype="auto"
)
model = PeftModel.from_pretrained(base_model, lora_model_path)
pipe = pipeline("text-generation", model=model, tokenizer=tokenizer)
print("✅ LLM 載入完成")

# === Embedding 查詢設定 ===
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model_dir = "./sustainability_embedding_model"
faiss_dir = "./faiss_data"
os.makedirs(faiss_dir, exist_ok=True)

embedding_tokenizer = AutoTokenizer.from_pretrained(model_dir)
embedding_base_model = AutoModel.from_pretrained(model_dir)

class EmbeddingModel(torch.nn.Module):
    def __init__(self, base_model, tokenizer, num_labels=3):
        super().__init__()
        self.model = base_model
        self.tokenizer = tokenizer
        self.classifier = torch.nn.Linear(base_model.config.hidden_size, num_labels)

    def get_embeddings(self, input_ids, attention_mask):
        outputs = self.model(input_ids=input_ids, attention_mask=attention_mask)
        token_embeddings = outputs.last_hidden_state
        extended_attention_mask = attention_mask.unsqueeze(-1)
        sum_embeddings = torch.sum(token_embeddings * extended_attention_mask, 1)
        sum_mask = torch.sum(extended_attention_mask, 1)
        sum_mask = torch.clamp(sum_mask, min=1e-9)
        return sum_embeddings / sum_mask

    def get_embedding(self, text, device):
        inputs = self.tokenizer(text, return_tensors='pt', padding='max_length', truncation=True, max_length=512).to(device)
        with torch.no_grad():
            return self.get_embeddings(inputs['input_ids'], inputs['attention_mask'])

embedding_model = EmbeddingModel(embedding_base_model, embedding_tokenizer)
embedding_model.load_state_dict(torch.load(os.path.join(model_dir, "model_weights.pt"), map_location=device), strict=False)
embedding_model = embedding_model.to(device).eval()

# === FAISS ===
faiss_index_path = os.path.join(faiss_dir, "faiss_index.index")
faiss_ids_path = os.path.join(faiss_dir, "faiss_ids.json")

if os.path.exists(faiss_index_path):
    try:
        index = faiss.read_index(faiss_index_path)
    except Exception as e:
        print("⚠ 讀取 FAISS 索引失敗：", e)
        index = None
else:
    index = None

if os.path.exists(faiss_ids_path):
    with open(faiss_ids_path, "r") as f:
        paragraph_ids = json.load(f)
else:
    paragraph_ids = []

# === Pydantic 模型定義 ===
class QueryRequest(BaseModel):
    message: str
    role: str

class UploadDocumentRequest(BaseModel):
    file_url: str
    file_name: str

class BatchUploadRequest(BaseModel):
    files: list[dict]  # [{"file_url": "...", "file_name": "..."}]

# === 文字清理和驗證函數（參考你提供的程式碼）===
def clean_text(text):
    """清理文字：去除多餘空白和無效字符"""
    if not text:
        return ""
    
    # 合併多個空白字符為單一空格
    text = re.sub(r'\s+', ' ', text)
    
    # 只保留：中文字、英文字母、數字、常見標點符號
    text = re.sub(r'[^\w\s\u4e00-\u9fff，。！？：；（）、""''「」『』《》〈〉.,:;!?"\'()[\]{}\-—%$€£¥₩₹]', '', text)
    
    return text.strip()

def is_valid_sentence(text):
    """檢查句子是否為有效句子（排除奇怪符號和無意義內容）"""
    if not text or len(text.strip()) < 10:  # 最少10個字符
        return False
    
    # 允許：中文字、中文標點、英文字母、阿拉伯數字、常見符號
    valid_pattern = re.compile(r'^[\u4e00-\u9fffA-Za-z0-9，。！？：；（）、「」『』《》〈〉.,:;!?"\'()\[\]{}\-—\s%$€£¥₩₹]{10,}$')
    
    if not valid_pattern.match(text):
        return False
    
    # 排除只有數字和符號的內容
    chinese_char_count = len(re.findall(r'[\u4e00-\u9fff]', text))
    english_char_count = len(re.findall(r'[A-Za-z]', text))
    
    # 確保至少有一定比例的中文或英文字符
    if (chinese_char_count + english_char_count) / len(text) < 0.3:
        return False
    
    # 排除頁碼、目錄等無意義內容
    meaningless_patterns = [
        r'^\d+\s*$',  # 純數字
        r'^第\s*\d+\s*頁\s*$',  # 頁碼
        r'^頁\s*\d+\s*$',  # 頁碼
        r'^目\s*錄\s*$',  # 目錄
        r'^\.\.\.\s*$',  # 省略號
        r'^[-—]+\s*$',  # 分隔線
        r'^\s*[IVX]+\s*\.\s*$',  # 羅馬數字章節
    ]
    
    for pattern in meaningless_patterns:
        if re.match(pattern, text.strip()):
            return False
    
    return True

def extract_meaningful_sentences(text):
    """從文字中提取有意義的句子"""
    if not text:
        return []
    
    meaningful_sentences = []
    
    # 先按段落分割
    paragraphs = text.split('\n\n')
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            continue
        
        # 按句號分割句子
        sentences = re.split(r'(?<=[。！？])\s*', paragraph)
        
        for sentence in sentences:
            if not sentence.strip():
                continue
            
            # 清理文字
            clean_sentence = clean_text(sentence)
            
            # 驗證是否為有效句子
            if is_valid_sentence(clean_sentence):
                meaningful_sentences.append(clean_sentence)
    
    return meaningful_sentences

# === 文件處理類別（優化版）===
class DocumentProcessor:
    def __init__(self):
        self.extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_doc,
            '.txt': self._extract_txt
        }
    
    def _extract_pdf(self, file_path):
        try:
            doc = fitz.open(file_path)
            text = ""
            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                if page_text.strip():
                    text += f"\n{page_text}\n"
            doc.close()
            
            # 如果 PyMuPDF 沒有提取到內容，嘗試使用 Tika
            if not text.strip():
                print("📄 PyMuPDF 未提取到內容，嘗試 Tika...")
                parsed = parser.from_file(file_path)
                text = parsed.get("content", "") or ""
            
            return text.strip()
        except Exception as e:
            print(f"PDF處理錯誤 {file_path}: {e}")
            # 備用方案：使用 Tika
            try:
                parsed = parser.from_file(file_path)
                return (parsed.get("content", "") or "").strip()
            except Exception as e2:
                print(f"Tika 備用方案也失敗: {e2}")
                return ""
    
    def _extract_docx(self, file_path):
        try:
            doc = docx.Document(file_path)
            text = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # 提取表格內容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))
            
            result = "\n".join(text)
            
            # 如果 python-docx 沒有提取到內容，嘗試使用 Tika
            if not result.strip():
                print("📄 python-docx 未提取到內容，嘗試 Tika...")
                parsed = parser.from_file(file_path)
                result = parsed.get("content", "") or ""
            
            return result.strip()
        except Exception as e:
            print(f"DOCX處理錯誤 {file_path}: {e}")
            # 備用方案：使用 Tika
            try:
                parsed = parser.from_file(file_path)
                return (parsed.get("content", "") or "").strip()
            except Exception as e2:
                print(f"Tika 備用方案也失敗: {e2}")
                return ""
    
    def _extract_doc(self, file_path):
        try:
            parsed = parser.from_file(file_path)
            content = parsed.get("content", "") or ""
            return content.strip()
        except Exception as e:
            print(f"DOC處理錯誤 {file_path}: {e}")
            return ""
    
    def _extract_txt(self, file_path):
        # 嘗試多種編碼
        encodings = ['utf-8', 'utf-8-sig', 'gbk', 'big5', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read().strip()
                    if content:
                        print(f"✅ 成功使用 {encoding} 編碼讀取 TXT 檔案")
                        return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"TXT處理錯誤 (encoding: {encoding}): {e}")
                continue
        
        print(f"⚠ 所有編碼都無法讀取 TXT 檔案: {file_path}")
        return ""
    
    def extract_text(self, file_path):
        if not os.path.exists(file_path):
            print(f"⚠ 檔案不存在: {file_path}")
            return ""
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            print(f"⚠ 檔案為空: {file_path}")
            return ""
        
        print(f"🔍 開始處理檔案: {file_path} (大小: {file_size} bytes)")
        
        ext = os.path.splitext(file_path)[1].lower()
        print(f"🔍 檔案副檔名: {ext}")
        
        if ext not in self.extractors:
            print(f"⚠ 不支援的檔案格式: {ext}")
            print(f"✅ 支援的格式: {list(self.extractors.keys())}")
            return ""
        
        extractor = self.extractors[ext]
        raw_text = extractor(file_path)
        
        if not raw_text or not raw_text.strip():
            print(f"⚠ 未提取到任何文字內容")
            return ""
        
        print(f"🔄 開始清理和驗證文字內容...")
        
        # 提取有意義的句子
        meaningful_sentences = extract_meaningful_sentences(raw_text)
        
        if not meaningful_sentences:
            print(f"⚠ 清理後沒有有效的句子")
            return ""
        
        # 重新組合有效句子
        clean_text_result = '\n'.join(meaningful_sentences)
        
        print(f"✅ 成功提取文字，原始長度: {len(raw_text)} 字元")
        print(f"✅ 清理後長度: {len(clean_text_result)} 字元")
        print(f"✅ 有效句子數量: {len(meaningful_sentences)}")
        print(f"🔍 預覽前3句: {meaningful_sentences[:3]}")
        
        return clean_text_result

# === 下載函數 ===
def download_file(url: str, local_path: str, timeout: int = 30) -> bool:
    """
    從 URL 下載檔案到本機路徑
    """
    try:
        print(f"🌐 開始下載檔案: {url}")
        
        # 設定請求標頭
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=timeout, stream=True)
        response.raise_for_status()
        
        # 檢查內容長度
        content_length = response.headers.get('Content-Length')
        if content_length:
            file_size = int(content_length)
            print(f"📊 檔案大小: {file_size} bytes")
            
            # 檢查檔案大小限制（例如 100MB）
            max_size = 100 * 1024 * 1024  # 100MB
            if file_size > max_size:
                print(f"⚠ 檔案過大: {file_size} bytes > {max_size} bytes")
                return False
        
        # 下載檔案
        with open(local_path, 'wb') as f:
            downloaded = 0
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
                    downloaded += len(chunk)
        
        print(f"✅ 檔案下載完成: {local_path} ({downloaded} bytes)")
        return True
        
    except requests.exceptions.RequestException as e:
        print(f"⚠ 網路請求錯誤: {e}")
        return False
    except Exception as e:
        print(f"⚠ 下載檔案時發生錯誤: {e}")
        return False

def rerank_documents(query: str, documents: list, top_k: int = 5):
    """
    使用 Cohere Rerank API 對文檔進行重新排序
    
    Args:
        query: 用戶查詢
        documents: 原始文檔列表（每個元素應該是包含 'text' 鍵的字典）
        top_k: 返回前 k 個最相關的文檔
    
    Returns:
        重新排序後的文檔列表
    """
    if not cohere_client:
        print("⚠️ Cohere 未初始化，跳過 Rerank，直接返回前 k 個文檔")
        return documents[:top_k]
    
    if not documents:
        print("⚠️ 沒有文檔需要 Rerank")
        return []
    
    try:
        print(f"🎯 開始 Rerank，候選文檔數：{len(documents)}，目標返回：{top_k} 個")
        
        # 準備文檔文字（Cohere 需要純文字列表）
        doc_texts = [doc.get('text', '') for doc in documents]
        
        # 過濾掉空文檔
        valid_docs = [(i, doc, text) for i, (doc, text) in enumerate(zip(documents, doc_texts)) if text.strip()]
        
        if not valid_docs:
            print("⚠️ 所有文檔都是空的，無法進行 Rerank")
            return []
        
        valid_indices, valid_documents, valid_texts = zip(*valid_docs)
        
        # 呼叫 Cohere Rerank API
        print(f"📡 呼叫 Cohere Rerank API...")
        rerank_response = cohere_client.rerank(
            model="rerank-multilingual-v3.0",  # 支援中文的模型
            query=query,
            documents=list(valid_texts),
            top_n=min(top_k, len(valid_texts)),  # 確保不超過可用文檔數
            return_documents=False  # 我們已經有原始文檔了
        )
        
        # 根據 Rerank 結果重新排序原始文檔
        reranked_documents = []
        for rank, result in enumerate(rerank_response.results):
            original_index = result.index
            reranked_doc = valid_documents[original_index].copy()
            reranked_doc['rerank_score'] = result.relevance_score
            reranked_doc['rerank_position'] = rank + 1
            reranked_documents.append(reranked_doc)
        
        print(f"✅ Rerank 完成，返回 {len(reranked_documents)} 個文檔")
        print(f"📊 Rerank 分數範圍：{reranked_documents[0]['rerank_score']:.3f} ~ {reranked_documents[-1]['rerank_score']:.3f}")
        
        return reranked_documents
        
    except Exception as e:
        print(f"❌ Rerank 失敗: {e}")
        print(f"⚠️ 降級為返回原始文檔的前 {top_k} 個")
        # 如果 Rerank 失敗，返回原始文檔
        return documents[:top_k]

# === 文字分塊函數（優化版）===
def chunk_text(text, max_length=400, overlap=50):
    """
    智能文字分塊，優先按句號分割，避免句子被截斷
    """
    if len(text) <= max_length:
        return [text]
    
    chunks = []
    sentences = re.split(r'(?<=[。！？])\s*', text)
    
    current_chunk = ""
    
    for sentence in sentences:
        # 如果當前句子本身就超過最大長度，強制分割
        if len(sentence) > max_length:
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
            
            # 對長句子進行強制分割
            start = 0
            while start < len(sentence):
                end = start + max_length
                chunks.append(sentence[start:end])
                start = end - overlap
        else:
            # 檢查加入這個句子是否會超過長度限制
            if len(current_chunk) + len(sentence) > max_length:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += sentence
    
    # 添加最後一個塊
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # 過濾掉太短的塊
    chunks = [chunk for chunk in chunks if len(chunk.strip()) >= 20]
    
    return chunks

# === 向量化處理函數（修正版）===
def process_document_to_vectors(text, source_name, source_url=None):
    """
    將文字轉換為向量並存入 FAISS 和 MongoDB
    """
    global index, paragraph_ids
    
    try:
        # 檢查必要的組件
        if not embedding_model:
            print("❌ Embedding 模型未載入")
            return 0, 0, [], 0
        
        if not text or not text.strip():
            print("❌ 輸入文字為空")
            return 0, 0, [], 0
        
        # 使用優化的分塊函數
        paragraphs = chunk_text(text)
        print(f"📝 文字分塊完成，共 {len(paragraphs)} 個段落")
        
        if not paragraphs:
            print("❌ 文字分塊後沒有內容")
            return 0, 0, [], 0
        
        paragraph_labels = [
            {
                "source": source_name,
                "category": "remote_document" if source_url else "uploaded_document",
                "chunk_id": i,
                "source_url": source_url
            } 
            for i in range(len(paragraphs))
        ]
        
        embeddings, new_ids = [], []
        processed_count = 0

        # 生成向量並準備文檔
        for para, label in zip(paragraphs, paragraph_labels):
            if len(para.strip()) < 10:  # 跳過太短的段落
                continue
                
            try:
                # 嘗試生成向量
                emb = embedding_model.get_embedding(para, device).cpu().numpy()
                embeddings.append(emb)
                para_id = str(uuid.uuid4())
                new_ids.append(para_id)
                processed_count += 1
                
                doc = {
                    "_id": para_id,
                    "text": para,
                    "category": label["category"],
                    "source": label["source"],
                    "chunk_id": label["chunk_id"],
                    "meta": {
                        "source_file": label["source"],
                        "category": label["category"],
                        "length": len(para),
                        "chunk_index": label["chunk_id"]
                    }
                }
                
                # 如果有 URL，添加到 meta 中
                if source_url:
                    doc["meta"]["source_url"] = source_url
                
                # 插入 MongoDB
                if local_api:
                    try:
                        # 修正：移除多餘的 https://
                        api_url = local_api if local_api.startswith('http') else f"https://{local_api}"
                        res = requests.post(f"{api_url}/insert_doc", json=doc, timeout=10)
                        res.raise_for_status()
                        print(f"✅ 成功插入文檔片段到 MongoDB: {para_id}")
                    except requests.exceptions.Timeout:
                        print(f"⚠️ MongoDB 插入超時: {para_id}")
                    except requests.exceptions.RequestException as e:
                        print(f"⚠️ MongoDB 請求失敗: {e}")
                    except Exception as e:
                        print(f"⚠️ 插入 MongoDB 失敗: {e}")
                else:
                    print("⚠️ LOCAL_API 未設定，跳過 MongoDB 插入")
                    
            except torch.cuda.OutOfMemoryError:
                print(f"❌ GPU 記憶體不足，無法處理段落 {processed_count}")
                # 清理 GPU 記憶體
                torch.cuda.empty_cache()
                continue
            except Exception as e:
                print(f"❌ 處理段落 {processed_count} 時發生錯誤: {e}")
                continue

        # 檢查是否有成功處理的向量
        if not embeddings:
            print("❌ 沒有成功生成任何向量")
            return 0, 0, [], 0

        # 更新 FAISS 索引
        try:
            new_embeddings = np.vstack(embeddings)
            
            if index is None:
                index = faiss.IndexFlatL2(new_embeddings.shape[1])
                print("🔍 創建新的 FAISS 索引")
            
            index.add(new_embeddings)
            paragraph_ids.extend(new_ids)

            # 保存 FAISS 索引
            try:
                faiss.write_index(index, faiss_index_path)
                with open(faiss_ids_path, "w") as f:
                    json.dump(paragraph_ids, f)
                
                print(f"✅ FAISS 索引已更新，新增 {len(new_ids)} 個向量")
            except Exception as e:
                print(f"❌ 保存 FAISS 索引失敗: {e}")
        
        except Exception as e:
            print(f"❌ 更新 FAISS 索引失敗: {e}")
            return len(paragraphs), 0, paragraphs[:3], 0
        
        return len(paragraphs), new_embeddings.shape[1] if embeddings else 0, paragraphs[:3], len(new_ids)
        
    except Exception as e:
        print(f"❌ 向量化處理失敗: {e}")
        print(f"📋 詳細錯誤: {traceback.format_exc()}")
        return 0, 0, [], 0

# === 聊天端點 ===
@app.post("/chat")
def chat(req: QueryRequest):
    prompt = f"""你是一位 ESG 永續報告分析師，請直接針對以下問題清楚簡短地回答，不要產生額外問題，此外，請用單純的文字回答，不要有格式：\n問題：{req.message}\n回答："""
    result = pipe(prompt, max_new_tokens=200)[0]['generated_text']
    answer = result.split("回答：")[-1].split("問題：")[0].strip() if "回答：" in result else result.strip()
    return {"reply": answer}


@app.post("/rag_chat")
def rag_chat(req: QueryRequest):
    if index is None:
        return {"reply": "⚠ FAISS 索引尚未初始化，請先上傳文件。", "references": []}

    # === 取得使用者角色 ===
    role = getattr(req, "role", None)
    mes = getattr(req, "message", None)
    print("role" : {role})
    print("問題" : {message})

    if not role:
        role = "外部人士"
        print("⚠ 沒收到前端角色資訊，預設為外部人士。")
    else:
        print(f"👤 使用者角色: {role}")

    # === 第一階段: FAISS 向量檢索 ===
    query_vec = embedding_model.get_embedding(req.message, device).cpu().numpy()
    initial_k = 25
    D, I = index.search(query_vec, initial_k)

    print(f"🔍 FAISS 檢索: 取得 {initial_k} 個候選文檔")
    top_ids = [paragraph_ids[idx] for idx in I[0] if idx < len(paragraph_ids)]
    print("🧾 FAISS 找到的 MongoDB _id:", top_ids)

    if not top_ids:
        return {"reply": "⚠ 沒有找到相關的文件段落。", "references": []}

    # === 取得原始段落 ===
    try:
        api_url = local_api if local_api.startswith('http') else f"https://{local_api}"
        response = requests.post(f"{api_url}/get_docs", json={"ids": top_ids}, timeout=10)
        documents = response.json()
        if isinstance(documents, dict) and "error" in documents:
            return {"reply": f"⚠ 本機 API 回傳錯誤: {documents['error']}", "references": []}
        print(f"📚 取得 {len(documents)} 個候選文檔")
    except Exception as e:
        return {"reply": f"⚠ 無法從本機 API 獲取段落: {str(e)}", "references": []}

    if not documents:
        return {"reply": "⚠ 沒有找到任何相關段落。", "references": []}

    # === 第二階段: Rerank 精選 ===
    print(f"🎯 開始 Rerank, 候選文檔數: {len(documents)}")
    final_k = 5
    reranked_documents = rerank_documents(req.message, documents, top_k=final_k)

    # === 信心門檻過濾 ===
    confidence_threshold = 0.8
    filtered_docs = [
        doc for doc in reranked_documents
        if isinstance(doc.get("rerank_score"), (int, float)) and doc["rerank_score"] >= confidence_threshold
    ]

    if not filtered_docs:
        print(f"⚠ 所有文檔信心分數都低於 {confidence_threshold}，不生成回答。")
        return {
            "reply": "⚠ 找不到足夠信心的資料，請嘗試詢問其他問題。",
            "references": [],
            "rerank_enabled": cohere_client is not None,
            "faiss_candidates": initial_k,
            "documents_used": 0,
            "role": role,
            "confidence_threshold": confidence_threshold
        }

    print(f"✅ 通過信心門檻的文檔數: {len(filtered_docs)} / {len(reranked_documents)}")
    reranked_documents = filtered_docs

    print("\n" + "="*80)
    print("📄 通過信心門檻的文檔:")
    print("="*80)
    for i, doc in enumerate(reranked_documents, 1):
        rerank_score = doc.get('rerank_score', 'N/A')
        source = doc.get('source', 'unknown')
        text_preview = doc.get('text', '')[:150]
        print(f"\n【文檔 {i}】")
        print(f"  Rerank 分數: {rerank_score}")
        print(f"  來源: {source}")
        print(f"  內容預覽: {text_preview}...")
    print("="*80 + "\n")

    # === 第三階段: 建立 prompt ===
    context_snippets = [doc['text'] for doc in reranked_documents]

    if role == "內部人士":
        role_instruction = "你可以提供精確的數字與細節，回答需盡量具體。"
    else:
        role_instruction = "請避免提供任何精確數字，改用『約』『大約』『少於』『超過』等模糊表達。"

    prompt = (
        f"你是一位 ESG 永續報告分析師。{role_instruction}\n"
        f"請根據以下參考資料回答問題，若找不到答案請誠實說明。\n"
        f"請用單純的文字回答，不要有格式。\n\n"
        f"參考資料:\n{chr(10).join(context_snippets)}\n\n"
        f"問題: {req.message}\n回答:"
    )

    # === 第四階段: 語言模型生成回答 ===
    result = pipe(prompt, max_new_tokens=3000)[0]['generated_text']
    answer = result.split("回答:")[-1].split("問題:")[0].strip() if "回答:" in result else result.strip()

    return {
        "reply": answer,
        "references": reranked_documents,
        "rerank_enabled": cohere_client is not None,
        "faiss_candidates": initial_k,
        "documents_used": len(reranked_documents),
        "role": role,
        "confidence_threshold": confidence_threshold
    }




@app.post("/gpt_chat")
def gpt_chat(req: QueryRequest):
    """
    直接請 Groq API 回答問題（不使用 RAG / 檢索）
    """
    # 檢查 Groq API Key
    if not groq_api_key:
        return {"error": "⚠ GROQ_API_KEY 未設定"}

    try:
        completion = groq_client.chat.completions.create(
            model="openai/gpt-oss-120b",
            messages=[
                {
                    "role": "system",
                    "content": "你是一位專業的 ESG 永續報告分析師。請根據你的專業知識回答用戶的問題。如果問題超出你的知識範圍，請誠實說明。此外，請用單純的文字回答，不要有格式。"
                },
                {
                    "role": "user",
                    "content": req.message
                }
            ],
            temperature=0.7,
            max_completion_tokens=1000,
            top_p=1,
            reasoning_effort="medium",
            stream=False,
            stop=None
        )
        
        answer = completion.choices[0].message.content
        return {
            "reply": answer,
            "references": [],
            "note": "基於模型自身知識回答（未檢索任何文件）"
        }

    except Exception as e:
        return {"error": f"⚠ Groq API 調用失敗：{str(e)}"}



@app.post("/gpt_rag_chat")
def gpt_rag_chat(req: QueryRequest):  # ← 注意函數名稱
    """
    使用 RAG 檢索相關文件，然後透過 Groq API 回答問題（含 Rerank）
    """
    # 檢查 Groq API Key
    if not groq_api_key:
        return {"error": "⚠ GROQ_API_KEY 未設定", "references": []}
    
    # 檢查 FAISS 索引是否存在
    if index is None:
        return {"reply": "⚠ FAISS 索引尚未初始化，請先上傳文件。", "references": []}
    
    try:
        # === RAG 檢索階段 ===
        query_vec = embedding_model.get_embedding(req.message, device).cpu().numpy()
        
        # 🔍 第一階段：FAISS 向量檢索（快速但不夠精準）
        # 先取更多候選文檔，之後用 Rerank 精選
        initial_k = 25  # 候選文檔數量
        D, I = index.search(query_vec, initial_k)
        
        print(f"🔍 FAISS 檢索：取得 {initial_k} 個候選文檔")
        print("🔍 FAISS 找到的向量 ID 索引：", I)
        
        top_ids = [paragraph_ids[idx] for idx in I[0] if idx < len(paragraph_ids)]
        print("🧾 FAISS 找到的 MongoDB _id：", top_ids)

        if not top_ids:
            # 如果沒有找到相關文檔，直接使用 Groq 回答
            try:
                completion = groq_client.chat.completions.create(
                    model="openai/gpt-oss-120b",
                    messages=[
                        {
                            "role": "system",
                            "content": "你是一位專業的 ESG 永續報告分析師。請根據你的專業知識回答用戶的問題。如果問題超出你的知識範圍，請誠實說明。此外，請用單純的文字回答，不要有格式。"
                        },
                        {
                            "role": "user",
                            "content": req.message
                        }
                    ],
                    temperature=0.7,
                    max_completion_tokens=1000,
                    top_p=1,
                    reasoning_effort="medium",
                    stream=False,
                    stop=None
                )
                
                answer = completion.choices[0].message.content
                return {
                    "reply": answer, 
                    "references": [], 
                    "note": "基於一般知識回答（未找到相關文檔）",
                    "rerank_enabled": False
                }
            
            except Exception as e:
                return {"error": f"⚠ Groq API 調用失敗：{str(e)}", "references": []}

        # === 獲取相關文檔 ===
        try:
            api_url = local_api if local_api.startswith('http') else f"https://{local_api}"
            response = requests.post(f"{api_url}/get_docs", json={"ids": top_ids}, timeout=10)
            documents = response.json()
            
            if isinstance(documents, dict) and "error" in documents:
                return {"reply": f"⚠ 本機 API 回傳錯誤：{documents['error']}", "references": []}
            
            print(f"📚 取得 {len(documents)} 個候選文檔")
            
        except Exception as e:
            return {"error": f"⚠ 無法從本機 API 獲取段落：{str(e)}", "references": []}

        if not documents:
            return {"error": "⚠ 沒有找到任何相關段落。", "references": []}

        # 🎯 第二階段：Rerank（精準排序）
        print(f"🎯 開始 Rerank，候選文檔數：{len(documents)}")
        final_k = 5  # 最終要使用的文檔數量
        reranked_documents = rerank_documents(req.message, documents, top_k=final_k)
        
        if not reranked_documents:
            print("⚠️ Rerank 後沒有文檔，使用原始文檔")
            reranked_documents = documents[:final_k]
        
        print(f"✅ 將使用 {len(reranked_documents)} 個文檔生成回答")

        print("\n" + "="*80)
        print("📄 Rerank 後採用的文檔：")
        print("="*80)
        for i, doc in enumerate(reranked_documents, 1):
            rerank_score = doc.get('rerank_score', 'N/A')
            source = doc.get('source', 'unknown')
            text_preview = doc.get('text', '')[:150]  # 只顯示前 150 個字
    
            print(f"\n【文檔 {i}】")
            print(f"  Rerank 分數：{rerank_score}")
            print(f"  來源：{source}")
            print(f"  內容預覽：{text_preview}...")
            print(f"  完整長度：{len(doc.get('text', ''))} 字元")
        print("="*80 + "\n")
        
        # === 準備上下文並呼叫 Groq API ===
        context_snippets = [doc['text'] for doc in reranked_documents]
        context_text = "\n\n".join(context_snippets)
        
        # 限制上下文長度避免超過 token 限制
        if len(context_text) > 4000:
            context_text = context_text[:4000] + "..."
            print(f"⚠️ 上下文過長，已截斷至 4000 字元")
        
        try:
            completion = groq_client.chat.completions.create(
                model="openai/gpt-oss-120b",
                messages=[
                    {
                        "role": "system",
                        "content": """你是一位專業的 ESG 永續報告分析師。請根據提供的參考資料回答用戶的問題。

回答要求：
1. 主要依據提供的參考資料回答
2. 如果參考資料中沒有相關資訊，請明確說明
3. 回答要準確、專業且易懂
4. 如果可能，請引用具體的數據或事實
5. 回答中不得包含任何特殊換行符號（如換行）、HTML 標籤（如 <br>），而應直接輸出正常文字段落。"""
                    },
                    {
                        "role": "user", 
                        "content": f"""參考資料：
{context_text}

問題：{req.message}

請根據以上參考資料回答問題。"""
                    }
                ],
                temperature=0.7,
                max_completion_tokens=1000,
                top_p=1,
                reasoning_effort="medium",
                stream=False,
                stop=None
            )
            
            answer = completion.choices[0].message.content
            
            return {
                "reply": answer,
                "references": reranked_documents,
                "note": "基於文檔內容 + Rerank + Groq GPT 回答",
                "rerank_enabled": cohere_client is not None,
                "documents_used": len(reranked_documents)
            }
            
        except Exception as e:
            return {"error": f"⚠ Groq API 調用失敗：{str(e)}", "references": reranked_documents}
    
    except Exception as e:
        print(f"gpt_rag_chat 發生錯誤: {e}")
        import traceback
        print(f"📋 詳細錯誤: {traceback.format_exc()}")
        return {"error": f"⚠ 處理請求時發生錯誤：{str(e)}", "references": []}


# === 上傳端點（原有的檔案上傳功能）===
@app.post("/upload-document")
async def upload_document(file: UploadFile = File(...)):
    print(f"📤 收到上傳檔案: {file.filename}")
    
    upload_dir = "./upload"
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, file.filename)
   
    try:
        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)

        processor = DocumentProcessor()
        text = processor.extract_text(file_path)
        
        # 清理臨時檔案
        os.remove(file_path)
        
        if not text.strip():
            return {"error": "⚠ 檔案無法擷取文字內容"}

        # 處理向量化
        paragraphs_count, dimension, preview, new_faiss_added = process_document_to_vectors(
            text, file.filename
        )

        return {
            "message": "✅ 文件已向量化",
            "paragraphs": paragraphs_count,
            "dimension": dimension,
            "preview": preview,
            "new_faiss_added": new_faiss_added,
            "text_length": len(text)
        }
    
    except Exception as e:
        # 確保清理臨時檔案
        if os.path.exists(file_path):
            os.remove(file_path)
        return {"error": f"⚠ 檔案處理失敗: {str(e)}"}


@app.post("/upload_documents")
async def upload_documents(req: UploadDocumentRequest):
    """
    從 URL 下載文件並進行向量化處理
    """
    local_file_path = None
    
    try:
        print(f"📤 收到文件處理請求: {req.file_name} from {req.file_url}")

        # 驗證輸入參數
        if not req.file_name:
            return {"success": False, "error": "檔案名稱為空"}

        if not req.file_url or not req.file_url.startswith(('http://', 'https://')):
            return {"success": False, "error": "無效的檔案 URL"}

        # 檢查檔案副檔名
        ext = os.path.splitext(req.file_name)[1].lower()
        supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        if ext not in supported_formats:
            return {
                "success": False,
                "error": f"不支援的檔案格式 '{ext}'",
                "supported_formats": supported_formats
            }

        # 建立臨時目錄
        temp_dir = "./temp_downloads"
        try:
            os.makedirs(temp_dir, exist_ok=True)
        except Exception as e:
            return {"success": False, "error": f"無法建立臨時目錄: {str(e)}"}

        # 生成唯一的本機檔案路徑
        unique_id = str(uuid.uuid4())[:8]
        safe_filename = f"{unique_id}_{req.file_name}"
        local_file_path = os.path.join(temp_dir, safe_filename)

        # 下載檔案
        try:
            if not download_file(req.file_url, local_file_path):
                return {"success": False, "error": "檔案下載失敗"}
        except Exception as e:
            return {"success": False, "error": f"下載檔案時發生錯誤: {str(e)}"}

        # 檢查下載的檔案
        if not os.path.exists(local_file_path):
            return {"success": False, "error": "下載的檔案不存在"}
        
        if os.path.getsize(local_file_path) == 0:
            return {"success": False, "error": "下載的檔案為空"}

        # 處理文件
        try:
            processor = DocumentProcessor()
            text = processor.extract_text(local_file_path)
        except Exception as e:
            return {
                "success": False, 
                "error": f"文件內容提取失敗: {str(e)}",
                "file_name": req.file_name
            }

        if not text or not text.strip():
            return {
                "success": False,
                "error": "檔案無法擷取文字內容",
                "file_name": req.file_name,
                "possible_reasons": [
                    "檔案可能是掃描的圖片 PDF（需要 OCR）",
                    "檔案內容為空或損壞",
                    "檔案格式不正確",
                    "檔案使用不支援的編碼",
                    "檔案包含太多無意義的內容（已被過濾）"
                ]
            }

        # 處理向量化
        try:
            paragraphs_count, dimension, preview, new_faiss_added = process_document_to_vectors(
                text, req.file_name, req.file_url
            )
            
            if paragraphs_count == 0:
                return {
                    "success": False,
                    "error": "向量化處理失敗，沒有生成任何向量",
                    "file_name": req.file_name
                }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"向量化處理失敗: {str(e)}",
                "file_name": req.file_name
            }

        return {
            "success": True,
            "message": "✅ 文件已成功處理並向量化",
            "file_name": req.file_name,
            "source_url": req.file_url,
            "paragraphs": paragraphs_count,
            "dimension": dimension,
            "preview": preview,
            "new_faiss_added": new_faiss_added,
            "text_length": len(text),
            "chunks_processed": paragraphs_count
        }

    except Exception as e:
        print(f"❌ upload_documents 處理失敗: {e}")
        print(f"📋 詳細錯誤: {traceback.format_exc()}")
        return {
            "success": False,
            "error": f"檔案處理失敗: {str(e)}",
            "file_name": getattr(req, 'file_name', 'unknown')
        }

    finally:
        # 確保清理臨時檔案
        if local_file_path and os.path.exists(local_file_path):
            try:
                os.remove(local_file_path)
                print(f"🗑️ 已清理臨時檔案: {local_file_path}")
            except Exception as e:
                print(f"⚠️ 清理臨時檔案失敗: {e}")

# === 批次處理端點（修正版）===
@app.post("/upload_documents_batch")
async def upload_documents_batch(req: BatchUploadRequest):
    """
    批次處理多個文件
    """
    try:
        # 驗證輸入
        if not req.files:
            return {"success": False, "error": "沒有提供任何檔案"}
        
        if len(req.files) > 50:  # 限制批次處理數量
            return {"success": False, "error": "批次處理檔案數量超過限制（最多 50 個）"}
        
        results = []
        success_count = 0
        error_count = 0

        print(f"🚀 開始批次處理 {len(req.files)} 個檔案")

        for index, file_info in enumerate(req.files):
            try:
                print(f"📄 處理檔案 {index + 1}/{len(req.files)}: {file_info.get('file_name', 'unknown')}")
                
                # 驗證單個檔案資訊
                if not isinstance(file_info, dict):
                    error_count += 1
                    results.append({
                        "file_index": index,
                        "file_name": "unknown",
                        "status": "error",
                        "result": {"success": False, "error": "檔案資訊格式錯誤"}
                    })
                    continue
                
                if 'file_name' not in file_info or 'file_url' not in file_info:
                    error_count += 1
                    results.append({
                        "file_index": index,
                        "file_name": file_info.get('file_name', 'unknown'),
                        "status": "error",
                        "result": {"success": False, "error": "缺少必要的檔案資訊（file_name 或 file_url）"}
                    })
                    continue
                
                # 建立單個檔案處理請求
                file_req = UploadDocumentRequest(**file_info)
                result = await upload_documents(file_req)
                
                # 根據結果更新計數器
                if result.get("success", False):
                    success_count += 1
                    status = "success"
                else:
                    error_count += 1
                    status = "error"
                
                results.append({
                    "file_index": index,
                    "file_name": file_info.get("file_name"),
                    "status": status,
                    "result": result
                })
                
            except Exception as e:
                error_count += 1
                print(f"❌ 批次處理單一檔案失敗: {e}")
                results.append({
                    "file_index": index,
                    "file_name": file_info.get("file_name", "unknown"),
                    "status": "error",
                    "result": {"success": False, "error": f"處理失敗: {str(e)}"}
                })

        # 計算處理統計
        total_files = len(req.files)
        success_rate = (success_count / total_files) * 100 if total_files > 0 else 0

        print(f"📊 批次處理完成: {success_count} 成功, {error_count} 失敗")

        return {
            "success": True,
            "message": f"批次處理完成",
            "total_files": total_files,
            "success_count": success_count,
            "error_count": error_count,
            "success_rate": round(success_rate, 2),
            "results": results
        }
        
    except Exception as e:
        print(f"❌ 批次上傳處理失敗: {e}")
        print(f"📋 詳細錯誤: {traceback.format_exc()}")
        return {
            "success": False, 
            "error": f"批次處理失敗: {str(e)}",
            "total_files": len(req.files) if hasattr(req, 'files') and req.files else 0,
            "results": []
        }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)
